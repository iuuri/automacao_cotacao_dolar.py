############# IMPORTAÇÃO DE BIBLIOTECAS #############
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.chrome.options import Options
from time import sleep
from docx import Document
from datetime import datetime
from docx.shared import Inches
from docx2pdf import convert
import os

#################################################### CRIAÇÃO DAS FUNÇÕES ####################################################

############# CONFIGURAR E ABRIR NAVEGADOR #############
def abrir_navegador():
    # Configurando navegador Chrome
    chrome_options = Options()
    chrome_options.add_argument("--headless")  # Rodar o navegador em modo headless 
    chrome_options.add_argument("--log-level=3")  # Desativar logs do ChromeDriver
    chrome_options.add_argument("--silent")
    chrome_options.add_experimental_option('excludeSwitches', ['enable-logging'])

    driver = webdriver.Chrome(options=chrome_options)

    try:
        print('Acessando site de cotação do dólar...')
        driver.get('https://wise.com/br/currency-converter/dolar-hoje')
        driver.execute_script("window.scrollBy(0, 220);")
        # Verifica se o elemento que indica que a página carregou está presente
        WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, "//div[@id = 'calculator']//span[@class = 'text-success']"))
        )
        print("Site acessado com sucesso.")
    except TimeoutException:
        print("Erro ao acessar o site. Verifique sua conexão com a internet ou se o site está no ar.")
        driver.quit()
        return None

    sleep(3)
    return driver

############# REGISTRAR COTAÇÃO ATUAL DO DOLAR #############
def extraindo_cotacao(driver):
    print('Extraindo cotação atual do dólar...')
    cotacao = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.XPATH, "//div[@id = 'calculator']//span[@class = 'text-success']"))
    ).text
    print('Extração concluída...')
    return cotacao

############# TIRAR PRINT DA TELA #############
def tirar_print(driver):
    print('Tirando print da tela...')
    elemento_print = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.XPATH, "//main[@id = 'main']//section"))
    )
    caminho_imagem = 'print_site_dolar.png'
    elemento_print.screenshot(caminho_imagem)
    return caminho_imagem

############# REGISTRAR DATA/HORA ATUAL #############
def registrar_data():
    print('Registrando data e hora da extração...')
    data_atual = datetime.now()
    data_formatada = data_atual.strftime('%d/%m/%Y')
    hora_formatada = data_atual.strftime('%H:%M')
    return data_formatada, hora_formatada

############# CRIAR DOCUMENTO WORD #############
def criar_docx(cotacao, hora_formatada, data_formatada, caminho_imagem):
    print('Criando arquivo docx...')
    arquivo_docx = Document()
    arquivo_docx.add_heading('Cotação atual do dólar')
    arquivo_docx.add_paragraph(f'Cotação atual do dólar {cotacao} no dia {data_formatada} às {hora_formatada}.\n'
                               'Print retirado do site Wise.')
    arquivo_docx.add_picture(caminho_imagem, width=Inches(6))
    arquivo_docx.add_paragraph('Automação criada por Iuri Souza')
    arquivo_docx.save('cotacao_dolar.docx')
    print('Arquivo criado com sucesso...')

############# CONVERTER DOCX EM PDF #############
def converter_docx_para_pdf():
    print('Convertendo arquivo docx para pdf...')
    try:
        convert("cotacao_dolar.docx", "cotacao_dolar.pdf")
        print('Conversão concluída com sucesso!')
    except Exception as e:
        print(f'Erro ao converter docx para pdf: {e}')

driver = abrir_navegador()
if driver:
    cotacao = extraindo_cotacao(driver)
    data_formatada, hora_formatada = registrar_data()
    caminho_imagem = tirar_print(driver)
    criar_docx(cotacao, hora_formatada, data_formatada, caminho_imagem)
    converter_docx_para_pdf()
